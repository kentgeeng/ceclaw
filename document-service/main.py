import os
import uuid
import httpx
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from config import LLM_URL, LLM_MODEL, PORT
from knowledge import store_document, search_documents
from generators.pptx_gen import generate_pptx
from generators.docx_gen import generate_docx
from generators.xlsx_gen import generate_xlsx
from generators.pdf_gen import generate_pdf

OUTPUT_DIR = "/tmp/ceclaw-docs"
STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

app = FastAPI(title="CECLAW Document Service")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

@app.get("/", response_class=HTMLResponse)
async def ui():
    with open(os.path.join(STATIC_DIR, "index.html"), encoding="utf-8") as f:
        return f.read()


class GenerateRequest(BaseModel):
    topic: str
    pages: int = 5
    theme: str = "dark"


@app.get("/health")
async def health():
    return {"status": "ok", "service": "ceclaw-document-service"}


@app.post("/generate/pptx")
async def gen_pptx(req: GenerateRequest):
    path = f"{OUTPUT_DIR}/{uuid.uuid4()}.pptx"
    await generate_pptx(req.topic, req.pages, path, req.theme)
    return FileResponse(path, filename=f"{req.topic}.pptx",
                        media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation")


@app.post("/generate/docx")
async def gen_docx(req: GenerateRequest):
    path = f"{OUTPUT_DIR}/{uuid.uuid4()}.docx"
    await generate_docx(req.topic, path)
    return FileResponse(path, filename=f"{req.topic}.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.post("/generate/xlsx")
async def gen_xlsx(req: GenerateRequest):
    path = f"{OUTPUT_DIR}/{uuid.uuid4()}.xlsx"
    await generate_xlsx(req.topic, path)
    return FileResponse(path, filename=f"{req.topic}.xlsx",
                        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.post("/generate/pdf")
async def gen_pdf(req: GenerateRequest):
    path = f"{OUTPUT_DIR}/{uuid.uuid4()}.pdf"
    await generate_pdf(req.topic, req.pages, path, req.theme)
    return FileResponse(path, filename=f"{req.topic}.pdf", media_type="application/pdf")


@app.post("/analyze")
async def analyze(file: UploadFile = File(...)):
    content = await file.read()
    filename = file.filename or "unknown"
    ext = filename.lower().split(".")[-1]

    if ext == "pdf":
        from ocr import ocr_pdf
        text = ocr_pdf(content)
    elif ext in ["png", "jpg", "jpeg", "tiff", "bmp"]:
        from ocr import ocr_image
        text = ocr_image(content)
    elif ext in ["txt", "md"]:
        text = content.decode("utf-8", errors="ignore")
    elif ext == "docx":
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == "xlsx":
        import io, openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content))
        rows = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                rows.append("\t".join([str(c) if c else "" for c in row]))
        text = "\n".join(rows)
    else:
        text = content.decode("utf-8", errors="ignore")

    # LLM 摘要
    async with httpx.AsyncClient(timeout=60) as c:
        r = await c.post(LLM_URL, json={
            "model": LLM_MODEL,
            "messages": [{"role": "user",
                          "content": f"請用繁體中文摘要以下文件內容，條列重點：\n\n{text[:3000]}"}],
            "temperature": 0
        })
    summary = r.json()["choices"][0]["message"]["content"]

    # 存進 Qdrant
    chunks_count = await store_document(filename, text, ext)

    return JSONResponse({"filename": filename, "summary": summary,
                         "chunks_stored": chunks_count, "text_length": len(text)})


@app.post("/analyze/batch")
async def analyze_batch(files: list[UploadFile] = File(...)):
    all_texts = []
    filenames = []
    for file in files:
        content = await file.read()
        filename = file.filename or "unknown"
        ext = filename.lower().split(".")[-1]
        if ext == "pdf":
            from ocr import ocr_pdf
            text = ocr_pdf(content)
        elif ext in ["png","jpg","jpeg","tiff","bmp"]:
            from ocr import ocr_image
            text = ocr_image(content)
        elif ext in ["txt","md"]:
            text = content.decode("utf-8", errors="ignore")
        elif ext == "docx":
            import io
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif ext == "xlsx":
            import io, openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content))
            rows = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    rows.append("\t".join([str(c) if c else "" for c in row]))
            text = "\n".join(rows)
        else:
            text = content.decode("utf-8", errors="ignore")
        all_texts.append(f"=== {filename} ===\n{text}")
        filenames.append(filename)
        await store_document(filename, text, ext)

    combined = "\n\n".join(all_texts)
    async with httpx.AsyncClient(timeout=120) as c:
        r = await c.post(LLM_URL, json={
            "model": LLM_MODEL,
            "messages": [{"role": "user",
                          "content": f"請用繁體中文分析以下 {len(files)} 份文件，整合重點並條列各文件關聯：\n\n{combined[:6000]}"}],
            "temperature": 0
        })
    summary = r.json()["choices"][0]["message"]["content"]
    return JSONResponse({"filenames": filenames, "count": len(files), "summary": summary})


@app.post("/ocr")
async def ocr_endpoint(file: UploadFile = File(...)):
    content = await file.read()
    filename = file.filename or "unknown"
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        from ocr import ocr_pdf
        text = ocr_pdf(content)
    else:
        from ocr import ocr_image
        text = ocr_image(content)
    await store_document(filename, text, "ocr")
    return JSONResponse({"filename": filename, "text": text})


@app.get("/search")
async def search(q: str, limit: int = 5):
    results = await search_documents(q, limit)
    return JSONResponse({"results": results})


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=PORT)
