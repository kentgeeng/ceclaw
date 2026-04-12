import httpx
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import LLM_URL, LLM_MODEL

SEARXNG_URL = "http://127.0.0.1:2337/v1/search"


async def search_web(query: str, num: int = 6) -> list:
    try:
        async with httpx.AsyncClient(timeout=15) as c:
            r = await c.post(SEARXNG_URL, json={"query": query})
            results = r.json().get("data", {}).get("web", [])[:num]
            return [{"title": x["title"], "description": x.get("description", "")} for x in results]
    except Exception:
        return []


async def llm_generate_report(topic: str, search_results: list) -> dict:
    context = "\n".join([f"- {r['title']}: {r['description'][:200]}" for r in search_results])
    prompt = f"""你是資深研究分析師，撰寫高品質繁體中文報告。

主題：{topic}

網路最新參考資料：
{context}

請生成一份完整報告，規則：
- 標題專業有力
- 摘要 100-150 字，點出核心洞察
- 3-5 個章節，每章 150-250 字
- 內容要有具體數據、案例、趨勢分析
- 台灣本地化視角
- 禁止空話和套話

只輸出 JSON，格式：
{{
  "title": "報告標題",
  "summary": "執行摘要（100-150字）",
  "sections": [
    {{"heading": "章節標題", "content": "章節詳細內容（150-250字）"}},
    ...
  ],
  "conclusion": "結論與建議（50-80字）"
}}"""

    async with httpx.AsyncClient(timeout=90) as c:
        r = await c.post(LLM_URL, json={
            "model": LLM_MODEL,
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0,
            "max_tokens": 3000,
        })
    text = r.json()["choices"][0]["message"]["content"].strip()
    text = re.sub(r'<\|channel\|>.*?<channel\|>', '', text, flags=re.DOTALL)
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)
    start = text.find("{")
    end = text.rfind("}") + 1
    return json.loads(text[start:end])


def _add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2a2a3a')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _set_heading_color(heading, rgb):
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*rgb)


async def generate_docx(topic: str, output_path: str) -> str:
    search_results = await search_web(topic)
    data = await llm_generate_report(topic, search_results)

    doc = Document()

    # 頁面設定
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # 預設字體
    doc.styles['Normal'].font.name = 'Noto Serif CJK TC'
    doc.styles['Normal'].font.size = Pt(11)

    # 封面標題
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(data["title"])
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x2E)
    title_run.font.name = 'Noto Serif CJK TC'

    # 副標題
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"CECLAW 研究報告  ·  {topic}")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0xAA)
    sub_run.font.name = 'Noto Serif CJK TC'

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # 執行摘要
    exec_heading = doc.add_heading('執行摘要', level=1)
    _set_heading_color(exec_heading, (0x1A, 0x56, 0xDB))

    summary_p = doc.add_paragraph(data["summary"])
    summary_p.runs[0].font.size = Pt(11)
    summary_p.runs[0].font.name = 'Noto Serif CJK TC'
    summary_p.paragraph_format.space_after = Pt(12)

    _add_horizontal_rule(doc)

    # 各章節
    for i, section_data in enumerate(data.get("sections", [])):
        doc.add_paragraph()
        heading = doc.add_heading(f"{i+1}. {section_data['heading']}", level=2)
        _set_heading_color(heading, (0x0D, 0x0D, 0x2E))

        content_p = doc.add_paragraph(section_data["content"])
        content_p.runs[0].font.size = Pt(11)
        content_p.runs[0].font.name = 'Noto Serif CJK TC'
        content_p.paragraph_format.space_after = Pt(10)
        content_p.paragraph_format.line_spacing = Pt(20)

    # 結論
    if data.get("conclusion"):
        _add_horizontal_rule(doc)
        doc.add_paragraph()
        conc_heading = doc.add_heading('結論與建議', level=1)
        _set_heading_color(conc_heading, (0x1A, 0x56, 0xDB))
        conc_p = doc.add_paragraph(data["conclusion"])
        conc_p.runs[0].font.size = Pt(11)
        conc_p.runs[0].font.name = 'Noto Serif CJK TC'

    # 頁腳
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("CECLAW Enterprise AI  ·  ColdElectric  ·  機密文件")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0xAA)

    doc.save(output_path)
    return output_path
